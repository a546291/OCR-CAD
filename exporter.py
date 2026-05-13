import os
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from database import get_db

EXPORT_DIR = os.path.join(os.path.dirname(__file__), "output", "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)


import re

def export_to_word(question_ids: list[str], title: str = "試卷") -> str:
    """依照選取的題目 ID 順序，產生 Word 試卷，回傳檔案路徑"""
    conn = get_db()
    doc = Document()

    # ── 頁面設定 ─────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ── 標題 ─────────────────────────────────────────────────────────────────
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.size = Pt(18)
    heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph()  # 空行

    # ── 取得並排序題目 ────────────────────────────────────────────────────────
    questions = []
    for q_id in question_ids:
        row = conn.execute(
            "SELECT id, subject, type, stem, options, order_index FROM questions WHERE id = ?", (q_id,)
        ).fetchone()
        if row:
            questions.append(dict(row))

    # 先依照 order_index 排序，確保每一群組內的順序符合原先 PDF
    questions.sort(key=lambda x: x["order_index"])

    # ── 按題型分組 ────────────────────────────────────────────────────────────
    from collections import defaultdict
    type_groups: dict[str, list] = defaultdict(list)
    for q in questions:
        type_groups[q["type"]].append(q)

    # ── 輸出題目 (按題型群組) ──────────────────────────────────────────────────
    global_num = 1
    for q_type, qs in type_groups.items():
        # 題型標題
        t_para = doc.add_paragraph()
        t_run = t_para.add_run(f"■ {q_type}")
        t_run.bold = True
        t_run.font.size = Pt(14)
        t_para.paragraph_format.space_after = Pt(6)

        for q in qs:
            # 題號 + 題幹
            stem_para = doc.add_paragraph()
            num_run = stem_para.add_run(f"{global_num}. ")
            num_run.bold = True
            num_run.font.size = Pt(11)

            # 取得本題的所有圖片資訊
            img_rows = conn.execute(
                "SELECT image_path FROM question_images WHERE question_id = ?", (q["id"],)
            ).fetchall()
            img_dict = {os.path.basename(r["image_path"]): r["image_path"] for r in img_rows}

            def render_text_with_images(para, text, font_size=11):
                # 依據標記分割字串
                parts = re.split(r'(\[IMAGE:\s*.*?\])', text)
                for part in parts:
                    if not part: continue
                    m = re.match(r'\[IMAGE:\s*(.+?)\]', part)
                    if m:
                        img_name = m.group(1)
                        if img_name in img_dict and os.path.exists(img_dict[img_name]):
                            try:
                                run = para.add_run()
                                run.add_picture(img_dict[img_name], width=Cm(10))
                            except Exception:
                                pass
                    else:
                        run = para.add_run(part)
                        run.font.size = Pt(font_size)

            render_text_with_images(stem_para, q["stem"])
            stem_para.paragraph_format.space_after = Pt(4)

            # 選項
            options = json.loads(q["options"]) if isinstance(q["options"], str) else q["options"]
            for opt in options:
                opt_para = doc.add_paragraph(style="List Bullet")
                opt_para.paragraph_format.left_indent = Cm(1)
                render_text_with_images(opt_para, opt, font_size=10)

            doc.add_paragraph()  # 題目間空行
            global_num += 1

    conn.close()

    # ── 存檔 ─────────────────────────────────────────────────────────────────
    safe_title = title.replace("/", "_").replace("\\", "_")
    output_path = os.path.join(EXPORT_DIR, f"{safe_title}.docx")
    doc.save(output_path)
    return output_path
